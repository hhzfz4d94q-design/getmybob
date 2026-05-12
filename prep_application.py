#!/usr/bin/env python3
"""
Application Prep Tool
---------------------
Given a job in the database, calls Claude to:
  - Tailor Geetanjali's resume (rewrite summary + re-rank/rewrite bullets)
  - Draft a 3-paragraph cover letter
  - Draft a LinkedIn outreach message to the hiring manager

Outputs go to ./applications/<company>_<title>_<date>/ as .docx + .txt files,
and the apply URL opens in the browser.

Run modes:
  python3 prep_application.py --test            # Verify API key is set up
  python3 prep_application.py --list            # Show top-scoring fresh jobs
  python3 prep_application.py --job <ID>        # Prep for a specific job
  python3 prep_application.py                   # Interactive picker
"""

import argparse
import json
import os
import re
import sqlite3
import sys
import subprocess
import webbrowser
from datetime import datetime
from urllib.request import Request, urlopen
from urllib.error import HTTPError

ROOT = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(ROOT, "jobs.db")
RESUME_PATH = os.path.join(ROOT, "resume.json")
ENV_PATH = os.path.join(ROOT, ".env")
APPLICATIONS_DIR = os.path.join(ROOT, "applications")

ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"
MODEL = "claude-sonnet-4-6"


# --- Config / API key ---------------------------------------------------

def load_api_key():
    """Read ANTHROPIC_API_KEY from .env file or environment."""
    key = os.environ.get("ANTHROPIC_API_KEY")
    if key:
        return key
    if os.path.exists(ENV_PATH):
        with open(ENV_PATH) as f:
            for line in f:
                line = line.strip()
                if line.startswith("ANTHROPIC_API_KEY="):
                    return line.split("=", 1)[1].strip().strip('"').strip("'")
    return None


def call_claude(prompt, system=None, max_tokens=4000):
    """Call Claude API with a single user message. Returns response text."""
    key = load_api_key()
    if not key:
        raise RuntimeError(
            "No ANTHROPIC_API_KEY found. See SETUP_AI_KEY.md for setup steps."
        )
    body = {
        "model": MODEL,
        "max_tokens": max_tokens,
        "messages": [{"role": "user", "content": prompt}],
    }
    if system:
        body["system"] = system
    req = Request(
        ANTHROPIC_API_URL,
        data=json.dumps(body).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "x-api-key": key,
            "anthropic-version": "2023-06-01",
        },
    )
    try:
        with urlopen(req, timeout=60) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        return data["content"][0]["text"]
    except HTTPError as e:
        msg = e.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"Claude API error ({e.code}): {msg}")


# --- Resume / job loading -----------------------------------------------

def load_resume():
    with open(RESUME_PATH) as f:
        return json.load(f)


def load_job(job_id_or_fp):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    # Try fingerprint exact, then prefix match, then rowid
    row = conn.execute(
        "SELECT * FROM jobs WHERE fingerprint=? OR fingerprint LIKE ? LIMIT 1",
        (job_id_or_fp, f"{job_id_or_fp}%"),
    ).fetchone()
    if not row:
        try:
            rowid = int(job_id_or_fp)
            row = conn.execute("SELECT * FROM jobs WHERE rowid=?", (rowid,)).fetchone()
        except ValueError:
            pass
    conn.close()
    if not row:
        return None
    return dict(row)


def list_jobs(limit=20):
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        """SELECT rowid, fingerprint, company_name, title, location, score, posted_at, first_seen
           FROM jobs WHERE senior=1 AND remote=1
           AND last_seen >= datetime('now','-14 days')
           ORDER BY score DESC, last_seen DESC LIMIT ?""",
        (limit,),
    ).fetchall()
    conn.close()
    return rows


# --- Prompt construction ------------------------------------------------

SYSTEM_PROMPT = """You are a senior career coach specializing in healthcare-IT leadership placements (Director / Senior Director / Principal / VP level). You help Geetanjali Arora tailor her application materials to specific roles.

CRITICAL RULES:
1. Never invent claims, achievements, employers, dates, or numbers that are not in Geetanjali's resume. You may rephrase and re-rank existing content; you may not fabricate.
2. Match the language of the job description — if it says "patient outcomes" use that phrase, if it says "GTM" use that, if it says "platform" prefer that over "system".
3. Keep tone confident but not boastful. Senior leaders speak with calm authority.
4. Always return valid JSON exactly matching the requested schema. No prose outside the JSON.
"""


def build_tailoring_prompt(resume, job):
    return f"""I have a job opening and Geetanjali's resume. Tailor the resume content for this specific role.

# JOB
Company: {job['company_name']}
Title: {job['title']}
Location: {job['location']}
Description:
{job['description'][:3500]}

# CURRENT RESUME (JSON)
{json.dumps(resume, indent=2)}

# TASK
Return ONLY a JSON object with this exact schema:

{{
  "tailored_summary": "2-3 sentence professional summary targeted at this specific role. Lead with the most relevant credentials.",
  "tailored_experience": [
    {{
      "company": "Pfizer",
      "title": "Principal, Digital Transformation",
      "start": "October 2022",
      "end": "November 2025",
      "tailored_bullets": [
        "Re-ranked / rewritten bullet 1 — emphasize what matters for this job",
        "Bullet 2",
        "..."
      ]
    }},
    {{ ... for each company in resume.experience, in same order ... }}
  ],
  "tailored_skills": [
    "Top 10-12 skills that matter most for this role, drawn from resume.skills (you may merge synonyms but not invent)"
  ],
  "match_analysis": {{
    "strengths": "1-2 sentences on Geetanjali's strongest fit for this role",
    "gaps": "1-2 sentences on any apparent gaps to address in cover letter / interview"
  }},
  "cover_letter": "Three short paragraphs. Paragraph 1: open with why this specific company/role is interesting (be specific, reference something concrete from the job description). Paragraph 2: 2-3 specific accomplishments that map to the job. Paragraph 3: brief close + call to action. No 'Dear Hiring Manager' / signature — just the body.",
  "linkedin_outreach": "A 4-5 sentence LinkedIn message to the hiring manager (or a senior employee at the company if HM not known). Brief intro, specific reason for reaching out about this role, one accomplishment that maps, ask for a short conversation. Casual but professional tone."
}}

Return ONLY the JSON, nothing else.
"""


# --- DOCX generation ----------------------------------------------------

def write_resume_docx(resume, tailored, out_path):
    """Generate a clean one-page-ish executive resume."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip3 install python-docx --break-system-packages")

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    p = resume["personal"]
    # Header — name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(p["name"].upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # Contact
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(f"{p['location']} | {p['phone']} | {p['email']} | {p['linkedin']}")
    contact_run.font.size = Pt(9)

    def section_heading(text):
        h = doc.add_paragraph()
        run = h.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        # underline-ish
        return h

    # Summary
    section_heading("Professional Summary")
    doc.add_paragraph(tailored["tailored_summary"])

    # Skills
    section_heading("Core Competencies")
    skills_p = doc.add_paragraph(" | ".join(tailored["tailored_skills"]))
    skills_p.runs[0].font.size = Pt(10)

    # Experience
    section_heading("Professional Experience")
    # Match tailored experience to resume.experience order
    for tex, rex in zip(tailored["tailored_experience"], resume["experience"]):
        # Company + dates line
        comp_p = doc.add_paragraph()
        comp_p.paragraph_format.space_before = Pt(6)
        cr = comp_p.add_run(f"{rex['company']}, {rex['location']}")
        cr.bold = True
        cr.font.size = Pt(10.5)
        comp_p.add_run(f"\t{tex.get('start', rex['start'])} – {tex.get('end', rex['end'])}").italic = True

        # Title
        title_p = doc.add_paragraph()
        tr = title_p.add_run(tex.get("title", rex["title"]))
        tr.italic = True
        tr.font.size = Pt(10.5)
        title_p.paragraph_format.space_after = Pt(2)

        for b in tex["tailored_bullets"]:
            bp = doc.add_paragraph(b, style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)

        # Prior roles same company
        for prior in rex.get("prior_roles_same_company", []):
            sub_p = doc.add_paragraph()
            sr = sub_p.add_run(prior["title"])
            sr.italic = True
            sr.font.size = Pt(10)
            sub_p.add_run(f"   ({prior['start']} – {prior['end']})").font.size = Pt(9)
            for b in prior["bullets"]:
                bp = doc.add_paragraph(b, style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)

    # Education
    section_heading("Education")
    for ed in resume["education"]:
        doc.add_paragraph(f"{ed['degree']}, {ed['institution']}")

    # Professional development
    if resume.get("professional_development"):
        section_heading("Professional Development")
        doc.add_paragraph(" | ".join(resume["professional_development"]))

    doc.save(out_path)


def write_cover_letter_docx(resume, job, cover_text, out_path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = resume["personal"]
    # Header
    h = doc.add_paragraph()
    h.add_run(p["name"]).bold = True
    doc.add_paragraph(p["location"])
    doc.add_paragraph(p["phone"])
    doc.add_paragraph(p["email"])
    doc.add_paragraph(p["linkedin"])
    doc.add_paragraph()
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph(f"Hiring Team")
    doc.add_paragraph(job["company_name"])
    doc.add_paragraph()
    doc.add_paragraph(f"Re: {job['title']}").bold = True
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Team,")
    doc.add_paragraph()

    for para in cover_text.strip().split("\n\n"):
        doc.add_paragraph(para.strip())
        doc.add_paragraph()

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(p["name"])

    doc.save(out_path)


# --- Main workflow ------------------------------------------------------

def sanitize(name):
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", name)[:60]


def extract_json(text):
    """Claude sometimes wraps JSON in ```json ... ``` — strip that."""
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return json.loads(text)


def prep_for_job(job):
    print(f"\nPreparing application for:")
    print(f"  {job['title']}  @  {job['company_name']}")
    print(f"  Location: {job['location']}")
    print(f"  URL: {job['url']}\n")

    resume = load_resume()
    print("Asking Claude to tailor your resume + draft cover letter & LinkedIn message…")
    prompt = build_tailoring_prompt(resume, job)
    response = call_claude(prompt, system=SYSTEM_PROMPT, max_tokens=4000)

    try:
        tailored = extract_json(response)
    except json.JSONDecodeError as e:
        print(f"\nERROR: Could not parse AI response as JSON: {e}")
        print("Raw response:\n", response[:1000])
        return None

    # Output folder
    today = datetime.now().strftime("%Y-%m-%d")
    folder_name = f"{today}_{sanitize(job['company_name'])}_{sanitize(job['title'])}"
    out_dir = os.path.join(APPLICATIONS_DIR, folder_name)
    os.makedirs(out_dir, exist_ok=True)

    resume_path = os.path.join(out_dir, "tailored_resume.docx")
    write_resume_docx(resume, tailored, resume_path)
    print(f"  Tailored resume → {resume_path}")

    cover_path = os.path.join(out_dir, "cover_letter.docx")
    write_cover_letter_docx(resume, job, tailored["cover_letter"], cover_path)
    print(f"  Cover letter   → {cover_path}")

    li_path = os.path.join(out_dir, "linkedin_outreach.txt")
    with open(li_path, "w") as f:
        f.write(tailored["linkedin_outreach"])
    print(f"  LinkedIn intro → {li_path}")

    # Match analysis
    analysis_path = os.path.join(out_dir, "match_notes.md")
    with open(analysis_path, "w") as f:
        f.write(f"# Match Analysis — {job['title']} @ {job['company_name']}\n\n")
        f.write(f"**Apply URL:** {job['url']}\n\n")
        f.write(f"**Strengths to lead with:**\n{tailored['match_analysis']['strengths']}\n\n")
        f.write(f"**Gaps to address in cover letter / interview:**\n{tailored['match_analysis']['gaps']}\n")
    print(f"  Match notes   → {analysis_path}")

    # Mark in DB as prepped
    mark_prepped(job["fingerprint"])

    # Open apply URL in browser
    if job.get("url"):
        print(f"\nOpening apply page in your browser…")
        try:
            webbrowser.open(job["url"])
        except Exception:
            pass

    # Open the output folder
    try:
        subprocess.run(["open", out_dir], check=False)
    except Exception:
        pass

    return out_dir


def mark_prepped(fingerprint):
    conn = sqlite3.connect(DB_PATH)
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS applications (
            fingerprint TEXT PRIMARY KEY,
            prepped_at TEXT,
            applied_at TEXT,
            status TEXT DEFAULT 'prepped',
            notes TEXT,
            follow_up_at TEXT
        );
    """)
    conn.execute(
        "INSERT OR REPLACE INTO applications (fingerprint, prepped_at, status) VALUES (?, ?, COALESCE((SELECT status FROM applications WHERE fingerprint=?), 'prepped'))",
        (fingerprint, datetime.now().isoformat(), fingerprint),
    )
    conn.commit()
    conn.close()


def interactive_pick():
    rows = list_jobs(20)
    if not rows:
        print("No senior + remote jobs found in the last 14 days. Run ./run.sh first.")
        return None
    print("\nTop fresh jobs:\n")
    for i, r in enumerate(rows, 1):
        print(f"  [{i:>2}] score={r[5]:>3}  {r[3][:55]:<55}  @ {r[2][:25]:<25}")
    print()
    choice = input("Pick a number (or 'q' to quit): ").strip()
    if choice.lower() == "q":
        return None
    try:
        idx = int(choice) - 1
        return load_job(rows[idx][1])
    except (ValueError, IndexError):
        print("Invalid choice.")
        return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--test", action="store_true", help="Verify API key is set up")
    parser.add_argument("--list", action="store_true", help="List top fresh jobs")
    parser.add_argument("--job", help="Job fingerprint or rowid to prep for")
    args = parser.parse_args()

    if args.test:
        key = load_api_key()
        if not key:
            print("FAIL: No ANTHROPIC_API_KEY found. See SETUP_AI_KEY.md.")
            sys.exit(1)
        if not key.startswith("sk-ant-"):
            print(f"FAIL: Key doesn't look right (should start with sk-ant-): {key[:15]}…")
            sys.exit(1)
        # Tiny test call
        try:
            r = call_claude("Reply with just the word: OK", max_tokens=10)
            if "OK" in r.upper():
                print("OK: AI key is configured correctly.")
            else:
                print(f"Unexpected response: {r}")
        except Exception as e:
            print(f"FAIL: {e}")
            sys.exit(1)
        return

    if args.list:
        for r in list_jobs(30):
            print(f"  [{r[1][:8]}] score={r[5]:>3}  {r[3]}  @ {r[2]}")
        return

    if args.job:
        job = load_job(args.job)
        if not job:
            print(f"Job not found: {args.job}")
            sys.exit(1)
        prep_for_job(job)
        return

    # Interactive
    job = interactive_pick()
    if job:
        prep_for_job(job)


if __name__ == "__main__":
    main()
